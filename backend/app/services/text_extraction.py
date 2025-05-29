import tempfile
import os
import logging
from typing import Optional, Tuple, Dict, Any

# For PDF extraction
try:
    from pdfminer.high_level import extract_text as pdf_extract_text
except ImportError:
    pdf_extract_text = None

# For PyMuPDF fallback
try:
    import fitz  # PyMuPDF
except ImportError:
    fitz = None

# For DOCX extraction
try:
    import docx
except ImportError:
    docx = None

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

class TextExtractionService:
    """Service for extracting text from different document formats"""
    
    def extract_text(self, file_path: str) -> Optional[str]:
        """
        Extract text from a file based on its extension
        
        Args:
            file_path: Path to the file
            
        Returns:
            Extracted text or None if extraction failed
        """
        _, file_extension = os.path.splitext(file_path)
        file_extension = file_extension.lower()
        
        if file_extension == '.pdf':
            return self._extract_from_pdf(file_path)
        elif file_extension == '.docx':
            return self._extract_from_docx(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")
    
    def _extract_from_pdf(self, file_path: str) -> Optional[str]:
        """Extract text from PDF file with fallback methods"""
        if pdf_extract_text is None:
            raise ImportError("pdfminer.six is not installed. Install it with 'pip install pdfminer.six'")
        
        # Try primary extraction method
        try:
            text = pdf_extract_text(file_path)
            if text and len(text.strip()) > 0:
                return text
            logger.warning(f"Primary PDF extraction returned empty text for {file_path}")
        except Exception as e:
            logger.warning(f"Primary PDF extraction failed for {file_path}: {e}")

        # Try PyMuPDF fallback if available
        if fitz is not None:
            try:
                logger.info(f"Trying PyMuPDF fallback for {file_path}")
                doc = fitz.open(file_path)
                text = ""
                for page in doc:
                    text += page.get_text()
                doc.close()
                if text and len(text.strip()) > 0:
                    logger.info(f"Successfully extracted text using PyMuPDF fallback for {file_path}")
                    return text
                logger.warning(f"PyMuPDF fallback returned empty text for {file_path}")
            except Exception as e:
                logger.warning(f"PyMuPDF fallback failed for {file_path}: {e}")
        
        # Try command-line tools if available
        try:
            import subprocess
            import platform
            
            # Check if pdftotext is available
            check_cmd = "where pdftotext" if platform.system() == "Windows" else "which pdftotext"
            subprocess.check_output(check_cmd, shell=True)
            
            # Create a temporary file for the text output
            with tempfile.NamedTemporaryFile(delete=False, suffix='.txt') as temp_file:
                temp_output_path = temp_file.name
            
            # Run pdftotext command
            logger.info(f"Trying pdftotext command line for {file_path}")
            cmd = f"pdftotext -layout \"{file_path}\" \"{temp_output_path}\""
            subprocess.run(cmd, shell=True, check=True)
            
            # Read the extracted text
            with open(temp_output_path, 'r', encoding='utf-8', errors='ignore') as f:
                text = f.read()
            
            # Clean up
            os.unlink(temp_output_path)
            
            if text and len(text.strip()) > 0:
                logger.info(f"Successfully extracted text using pdftotext for {file_path}")
                return text
            logger.warning(f"pdftotext returned empty text for {file_path}")
        except (subprocess.SubprocessError, FileNotFoundError, ImportError):
            # Command line tool not available, just continue
            pass
        except Exception as e:
            logger.warning(f"Command line extraction failed for {file_path}: {e}")
        
        logger.error(f"All PDF extraction methods failed for {file_path}")
        return None
    
    def _extract_from_docx(self, file_path: str) -> Optional[str]:
        """Extract text from DOCX file"""
        if docx is None:
            raise ImportError("python-docx is not installed. Install it with 'pip install python-docx'")
        
        try:
            doc = docx.Document(file_path)
            full_text = []
            
            # Extract text from paragraphs
            for para in doc.paragraphs:
                full_text.append(para.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        full_text.append(cell.text)
            
            return '\n'.join(full_text)
        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {e}")
            return None
    
    async def extract_text_from_upload(self, file_content: bytes, filename: str) -> Tuple[bool, str, Optional[Dict[str, Any]]]:
        """
        Extract text from an uploaded file with metadata
        
        Args:
            file_content: Binary content of the file
            filename: Name of the file with extension
            
        Returns:
            Tuple containing:
            - Success flag (bool)
            - Extracted text or error message (str)
            - Metadata dictionary (Dict) or None if extraction failed
        """
        # Create a temporary file
        with tempfile.NamedTemporaryFile(delete=False) as temp_file:
            temp_file.write(file_content)
            temp_file_path = temp_file.name
        
        try:
            # Extract text from the temporary file
            file_size = os.path.getsize(temp_file_path) / (1024 * 1024)  # Convert to MB
            text = self.extract_text(temp_file_path)
            
            if text and len(text.strip()) > 0:
                # Create metadata
                metadata = {
                    "file_name": filename,
                    "file_size_mb": round(file_size, 2),
                    "text_length": len(text),
                    "pages": self._estimate_page_count(text),
                    "extraction_method": "standard",
                    "extraction_status": "success",
                    "original_filename": filename
                }
                return True, text, metadata
            else:
                return False, "No text could be extracted from the file", None
        except Exception as e:
            return False, f"Error extracting text: {str(e)}", None
        finally:
            # Clean up the temporary file
            os.unlink(temp_file_path)
    
    def _estimate_page_count(self, text: str) -> int:
        """
        Estimate the number of pages based on text length
        
        Args:
            text: The extracted text
            
        Returns:
            Estimated number of pages
        """
        # Very rough estimate: ~3000 characters per page on average
        char_per_page = 3000
        return max(1, round(len(text) / char_per_page)) 